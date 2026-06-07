"""
Lector de documentos .docx. Extraído de engine.py.
"""
from __future__ import annotations

import logging
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from core.ocr.document_model import Document

logger = logging.getLogger(__name__)

try:
    import docx as _docx
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# Estilos de python-docx que corresponden a encabezados (Heading 1…6)
_HEADING_STYLES = {f"Heading {i}" for i in range(1, 7)}


def read_docx(docx_path: str) -> str:
    """Lee un archivo .docx y devuelve el texto como string (API original)."""
    if not _DOCX_OK:
        return "Error: python-docx no disponible. Instalar: pip install python-docx"
    try:
        doc = _docx.Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX error: {e}")
        return f"Error leyendo Word: {e}"


def read_docx_document(docx_path: str) -> Document:
    """
    Lee un .docx y devuelve un Document estructurado con TextBlocks tipados.
    Los encabezados (Heading 1-6) se mapean a BlockType.HEADING.
    Los ítems de lista se mapean a BlockType.LIST_ITEM.
    """
    from core.ocr.document_model import BlockType, Document, DocumentPage, TextBlock

    doc_obj = Document(source_path=str(docx_path), source_type="docx")

    if not _DOCX_OK:
        return doc_obj

    try:
        word_doc = _docx.Document(docx_path)
        page_doc = DocumentPage(page_number=1, source_path=str(docx_path))

        for para in word_doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            if style_name in _HEADING_STYLES:
                level = int(style_name.split()[-1])
                block = TextBlock(
                    text=text,
                    block_type=BlockType.HEADING,
                    heading_level=level,
                )
            elif "List" in style_name:
                block = TextBlock(text=text, block_type=BlockType.LIST_ITEM)
            else:
                block = TextBlock(text=text, block_type=BlockType.PARAGRAPH)

            page_doc.blocks.append(block)

        doc_obj.pages.append(page_doc)

    except Exception as exc:
        logger.error("read_docx_document: error en '%s': %s", docx_path, exc)

    return doc_obj
